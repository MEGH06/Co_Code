import docx
from docx import Document

def export_to_docx(content, output_path):
    """Export consolidated content to a Word document."""
    doc = Document()
    # Add summarized text
    doc.add_heading("Consolidated Notes", level=1)
    doc.add_paragraph(content["text"])
    # Add tables
    for table in content["tables"]:
        doc.add_heading("Table", level=2)
        table_obj = doc.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, cell in enumerate(row):
                table_obj.cell(i, j).text = cell
    # Add images
    for img_path in content["images"]:
        doc.add_picture(img_path, width=docx.shared.Inches(4))
    doc.save(output_path)