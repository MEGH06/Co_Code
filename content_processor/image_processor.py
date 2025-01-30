from transformers import CLIPProcessor, CLIPModel, pipeline
from PIL import Image
import docx
from docx import Document
import torch

# Load models
clip_model=CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
clip_processor=CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
summarizer=pipeline("summarization",model="facebook/bart-large-cnn")

def check_relevance(image:Image,text:str)->float:
    inputs=clip_processor(text=[text],images=image,return_tensors="pt",padding=True)
    with torch.no_grad():
        outputs=clip_model(**inputs)
    return outputs.logits_per_image.item() # Higher means more relevant

def create_summary(text:str)->str:
    summary=summarizer(text,max_length=150,min_length=50,do_sample=False)
    return summary[0]["summary_text"]

def generate_docx(images_text_list:list,output_file:str="output.docx"):
    doc=Document()
    for item in images_text_list:
        text=item.get("text","")
        images=item.get("images",[])
        summary=create_summary(text) if text else ""
        
        if summary:
            doc.add_paragraph(summary)
        
        for img_path in images:
            image=Image.open(img_path).convert("RGB")
            relevance_score=check_relevance(image,summary)
            if relevance_score>2.0: # Adjust threshold as needed
                doc.add_picture(img_path,width=docx.shared.Inches(4))
        
    doc.save(output_file)

# Example Usage
documents_data=[
    {"text":"This is some text about AI and deep learning.","images":["ai_image1.jpg","image2.png"]},
    {"text":"Climate change is affecting our planet.","images":["climate.jpg","tree.jpg"]}
]

generate_docx(documents_data,"final_output.docx")