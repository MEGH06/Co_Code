from docx import Document
import os
import zipfile
import pandas as pd

def extract_docx_content(file_path, output_dir="media"):
    """Extract text, tables, and images from a Word document."""
    content = {"text": [], "tables": [], "images": []}

    # Extract text and tables
    doc = Document(file_path)
    for para in doc.paragraphs:
        if para.text.strip():  # Skip empty paragraphs
            content["text"].append(para.text.strip())
    for table in doc.tables:
        table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        content["tables"].append(pd.DataFrame(table_data))  # Add table as DataFrame

    # Extract images
    os.makedirs(output_dir, exist_ok=True)
    with zipfile.ZipFile(file_path, 'r') as docx:
        media_files = [
            f for f in docx.namelist()
            if f.startswith('word/media/') and f.lower().endswith(('.jpg', '.png', '.jpeg'))
        ]
        for media in media_files:
            output_path = os.path.join(output_dir, os.path.basename(media))
            with open(output_path, 'wb') as img_file:
                img_file.write(docx.read(media))
            content["images"].append(output_path)

    return content

# Example usage
file_path = r"C:\Users\Kruttika\OneDrive\Desktop\Kruttika\sample.docx"
output_dir = "media"  # Directory to store extracted images
content = extract_docx_content(file_path, output_dir)

# Print results
print("Text Content:", content["text"])
print("\nTables:")
for i, table_df in enumerate(content["tables"]):
    print(f"\nTable {i + 1}:\n", table_df)
print("\nImages Extracted:", content["images"])


print(content)