import pandas as pd
import docx
from PIL import Image
from content_processor.image_processor import check_relevance  # Import the image relevance logic

def export_table_to_docx(table: pd.DataFrame, doc: docx.Document):
    """export a pandas DataFrame to a table in a docx document."""
    table_in_docx = doc.add_table(rows=1, cols=len(table.columns))
    hdr_cells = table_in_docx.rows[0].cells
    for i, column_name in enumerate(table.columns):
        hdr_cells[i].text = column_name
    for _, row in table.iterrows():
        row_cells = table_in_docx.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

def export_images_and_text_to_docx(data_dict, output_file_path="../output/final_consolidated_notes.docx"):
    """used to generate a docx document with summarized text and relevant images from the input dictionary."""
    doc = docx.Document()
    for text in data_dict.get("text", []):
        doc.add_paragraph(text)
    for img_path in data_dict.get("images", []):
        image = Image.open(img_path).convert("RGB")
        for text in data_dict.get("text", []):
            relevance_score = check_relevance(image, text)
            if relevance_score > 2.0:  # Adjust threshold as needed
                doc.add_picture(img_path, width=docx.shared.Inches(4))
    doc.save(output_file_path)

def export_tables_to_docx(data_dict, output_file_path="../output/consolidated_tables.docx"):
    """exporting all tables from input into docx document."""
    doc = docx.Document()
    for table in data_dict.get("tables", []):
        export_table_to_docx(table, doc)
    doc.save(output_file_path)